"""문서 종류별 텍스트 추출 전략을 모아둔 모듈이다."""

from __future__ import annotations

import logging
import threading
from pathlib import Path
from typing import Any

from common.document_types import (
    ExtractedTextResult,
    MIN_DIRECT_TEXT_KO_EN_CHARS_FOR_SKIP_OCR,
    MIN_DIRECT_TEXT_LENGTH_FOR_SKIP_OCR,
    MIN_DIRECT_TEXT_SCORE_FOR_SKIP_OCR,
    MIN_MERGEABLE_TEXT_SCORE,
    MIN_OCR_SCORE_ADVANTAGE,
    OCR_LINE_GROUPING_TOLERANCE,
    OCR_RENDER_DPI,
    PDF_CLASSIFICATION_SAMPLE_PAGES,
    PDF_SCANNED_TEXT_SCORE_THRESHOLD,
    PDF_TEXT_DOCUMENT_AVG_SCORE_THRESHOLD,
    PDF_TEXT_DOCUMENT_MIN_TEXT_PAGES,
    PageExtractionResult,
    PageTextSnapshot,
    PdfDocumentSignal,
    PdfPageSignal,
)
from common.file_storage import resolve_absolute_path
from common.text_normalizer import (
    build_strategy,
    compute_quality_score,
    count_meaningful_chars,
    deduplicate_text_sections,
    infer_document_kind,
    is_llm_output_acceptable,
    normalize_document_pages,
    normalize_extracted_text,
    normalize_line,
    run_llm_normalization,
    score_extracted_text,
    should_run_llm_cleaning,
)

logger = logging.getLogger(__name__)
_OCR_ENGINE_LOCAL = threading.local()

'''
extract_text_from_file() 진입점
PDF 분류, direct/OCR/hybrid 추출
'''

def _extract_plain_text(abs_path: Path) -> str | None:
    return normalize_extracted_text(abs_path.read_text(encoding="utf-8", errors="ignore"))


def _extract_docx_text(abs_path: Path) -> str | None:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.info("python-docx is not installed. Skipping DOCX extraction for %s", abs_path)
        return None

    try:
        document = DocxDocument(abs_path)
    except Exception:
        logger.exception("Failed to open DOCX file for extraction: %s", abs_path)
        return None

    sections: list[str] = []

    for paragraph in document.paragraphs:
        line = normalize_line(paragraph.text)
        if line:
            sections.append(line)

    for table in document.tables:
        for row in table.rows:
            row_cells = [normalize_line(cell.text) for cell in row.cells]
            row_cells = [cell for cell in row_cells if cell]
            if row_cells:
                sections.append(" | ".join(row_cells))

    return normalize_extracted_text("\n".join(sections))


def _get_rapidocr_engine():
    engine = getattr(_OCR_ENGINE_LOCAL, "rapidocr_engine", None)
    if engine is not None:
        return engine

    try:
        from rapidocr_onnxruntime import RapidOCR
    except ImportError:
        logger.warning("rapidocr-onnxruntime is not installed. OCR fallback is disabled.")
        return None

    try:
        engine = RapidOCR()
        _OCR_ENGINE_LOCAL.rapidocr_engine = engine
        return engine
    except Exception:
        logger.exception("Failed to initialize RapidOCR engine.")
        return None


def _should_run_ocr(direct_text: str | None) -> bool:
    normalized = normalize_extracted_text(direct_text or "")
    if not normalized:
        return True

    direct_score = score_extracted_text(normalized)
    meaningful_chars = count_meaningful_chars(normalized)

    return not (
        len(normalized) >= MIN_DIRECT_TEXT_LENGTH_FOR_SKIP_OCR
        or meaningful_chars >= MIN_DIRECT_TEXT_KO_EN_CHARS_FOR_SKIP_OCR
        or direct_score >= MIN_DIRECT_TEXT_SCORE_FOR_SKIP_OCR
    )


def _sort_ocr_results(ocr_result: list[Any]) -> list[str]:
    sortable_items: list[tuple[float, float, str]] = []

    for item in ocr_result:
        if not isinstance(item, (list, tuple)) or len(item) < 2:
            continue

        box = item[0]
        text = item[1]
        if not isinstance(text, str) or not text.strip():
            continue
        if not isinstance(box, (list, tuple)) or not box:
            continue

        points: list[tuple[float, float]] = []
        for point in box:
            if (
                isinstance(point, (list, tuple))
                and len(point) >= 2
                and isinstance(point[0], (int, float))
                and isinstance(point[1], (int, float))
            ):
                points.append((float(point[0]), float(point[1])))

        if not points:
            continue

        top = min(y for _, y in points)
        left = min(x for x, _ in points)
        sortable_items.append((top, left, text.strip()))

    sortable_items.sort(key=lambda value: (value[0], value[1]))

    grouped_lines: list[list[tuple[float, str]]] = []
    current_top: float | None = None

    for top, left, text in sortable_items:
        if current_top is None or abs(top - current_top) > OCR_LINE_GROUPING_TOLERANCE:
            grouped_lines.append([(left, text)])
            current_top = top
            continue
        grouped_lines[-1].append((left, text))

    ordered_lines: list[str] = []
    for line in grouped_lines:
        line.sort(key=lambda value: value[0])
        ordered_lines.append(" ".join(text for _, text in line))

    return ordered_lines


def _flatten_rapidocr_result(raw_result: Any) -> list[list[Any]]:
    flattened: list[list[Any]] = []

    def is_point(value: Any) -> bool:
        return (
            isinstance(value, (list, tuple))
            and len(value) >= 2
            and isinstance(value[0], (int, float))
            and isinstance(value[1], (int, float))
        )

    def is_box(value: Any) -> bool:
        return (
            isinstance(value, (list, tuple))
            and len(value) >= 4
            and all(is_point(point) for point in value)
        )

    def append_if_valid(box: Any, text: Any) -> None:
        if not is_box(box) or not isinstance(text, str):
            return
        normalized_text = text.strip()
        if not normalized_text:
            return
        flattened.append([box, normalized_text])

    def walk(node: Any) -> None:
        if node is None:
            return

        if isinstance(node, dict):
            possible_box = (
                node.get("box")
                or node.get("bbox")
                or node.get("points")
                or node.get("polygon")
                or node.get("dt_polys")
            )
            possible_text = (
                node.get("text")
                or node.get("transcription")
                or node.get("rec_text")
                or node.get("label")
            )
            append_if_valid(possible_box, possible_text)
            for value in node.values():
                walk(value)
            return

        if isinstance(node, (list, tuple)):
            if len(node) >= 2 and is_box(node[0]) and isinstance(node[1], str):
                append_if_valid(node[0], node[1])
                return

            if (
                len(node) >= 2
                and is_box(node[0])
                and isinstance(node[1], (list, tuple))
                and len(node[1]) >= 1
                and isinstance(node[1][0], str)
            ):
                append_if_valid(node[0], node[1][0])
                return

            for item in node:
                walk(item)

    walk(raw_result)

    deduped: list[list[Any]] = []
    seen: set[tuple[str, tuple[tuple[float, float], ...]]] = set()

    for box, text in flattened:
        box_key = tuple((float(x), float(y)) for x, y in box)
        key = (text, box_key)
        if key in seen:
            continue
        seen.add(key)
        deduped.append([box, text])

    return deduped


def _extract_ocr_text_from_page(page, abs_path: Path) -> str | None:
    rapid_ocr = _get_rapidocr_engine()
    if rapid_ocr is None:
        return None

    tmp_path: Path | None = None

    try:
        import fitz

        matrix = fitz.Matrix(OCR_RENDER_DPI / 72, OCR_RENDER_DPI / 72)
        pixmap = page.get_pixmap(matrix=matrix, alpha=False)
        tmp_path = abs_path.with_suffix(f".page{page.number + 1}.png")
        pixmap.save(tmp_path)

        raw_result, _ = rapid_ocr(str(tmp_path))
        normalized_items = _flatten_rapidocr_result(raw_result)

        if not normalized_items:
            return None

        ordered_lines = _sort_ocr_results(normalized_items)
        return normalize_extracted_text("\n".join(ordered_lines))
    except Exception:
        logger.exception(
            "RapidOCR extraction failed for file=%s page=%s dpi=%s",
            abs_path,
            page.number + 1,
            OCR_RENDER_DPI,
        )
        return None
    finally:
        if tmp_path is not None:
            try:
                tmp_path.unlink(missing_ok=True)
            except OSError:
                logger.warning("Failed to remove temporary OCR image file=%s", tmp_path)


def _extract_text_blocks_from_page(page) -> str | None:
    blocks = page.get_text("blocks", sort=True) or []
    block_texts: list[str] = []

    for block in blocks:
        if not isinstance(block, (list, tuple)) or len(block) < 5:
            continue
        block_text = normalize_extracted_text(str(block[4]))
        if block_text:
            block_texts.append(block_text)

    return normalize_extracted_text("\n\n".join(block_texts))


def _build_pdf_page_signal(page) -> PdfPageSignal:
    direct_text = _extract_text_blocks_from_page(page)
    return PdfPageSignal(
        page_number=page.number + 1,
        direct_text=direct_text,
        direct_score=score_extracted_text(direct_text),
        meaningful_chars=count_meaningful_chars(direct_text),
        image_count=len(page.get_images(full=True)),
        block_count=len(page.get_text("blocks", sort=True) or []),
    )


def _classify_pdf_document(document) -> PdfDocumentSignal:
    sample_pages = min(document.page_count, PDF_CLASSIFICATION_SAMPLE_PAGES)
    page_signals = [_build_pdf_page_signal(document[index]) for index in range(sample_pages)]

    text_pages = sum(
        1
        for signal in page_signals
        if signal.direct_score >= PDF_TEXT_DOCUMENT_AVG_SCORE_THRESHOLD
    )
    low_text_pages = sum(
        1
        for signal in page_signals
        if signal.direct_score <= PDF_SCANNED_TEXT_SCORE_THRESHOLD
    )
    average_score = (
        sum(signal.direct_score for signal in page_signals) / len(page_signals)
        if page_signals
        else 0
    )

    if text_pages >= min(PDF_TEXT_DOCUMENT_MIN_TEXT_PAGES, sample_pages) or average_score >= 100:
        source_type = "digital_pdf_text"
    elif low_text_pages == sample_pages:
        source_type = "scanned_pdf_ocr"
    else:
        source_type = "mixed_pdf"

    return PdfDocumentSignal(source_type=source_type, page_signals=page_signals)


def _select_page_text(*, direct_text: str | None, ocr_text: str | None) -> PageExtractionResult:
    direct_score = score_extracted_text(direct_text)
    ocr_score = score_extracted_text(ocr_text)

    if direct_score == 0 and ocr_score == 0:
        return PageExtractionResult(None, direct_text, ocr_text, direct_score, ocr_score, False)

    if direct_score >= MIN_DIRECT_TEXT_SCORE_FOR_SKIP_OCR:
        return PageExtractionResult(direct_text, direct_text, ocr_text, direct_score, ocr_score, False)

    if direct_score >= MIN_MERGEABLE_TEXT_SCORE and ocr_score >= MIN_MERGEABLE_TEXT_SCORE:
        merged_text = deduplicate_text_sections([direct_text or "", ocr_text or ""])
        merged_score = score_extracted_text(merged_text)
        if merged_score >= max(direct_score, ocr_score) + 10:
            return PageExtractionResult(merged_text, direct_text, ocr_text, direct_score, ocr_score, True)

    if ocr_score >= direct_score + MIN_OCR_SCORE_ADVANTAGE:
        return PageExtractionResult(ocr_text, direct_text, ocr_text, direct_score, ocr_score, True)

    return PageExtractionResult(direct_text, direct_text, ocr_text, direct_score, ocr_score, ocr_score > 0)


def _extract_text_pdf_page(page) -> str | None:
    return _extract_text_blocks_from_page(page)


def _extract_mixed_pdf_page(page, abs_path: Path) -> PageExtractionResult:
    direct_text = _extract_text_pdf_page(page)
    ocr_text = _extract_ocr_text_from_page(page, abs_path) if _should_run_ocr(direct_text) else None
    return _select_page_text(direct_text=direct_text, ocr_text=ocr_text)


def _extract_scanned_pdf_page(page, abs_path: Path) -> PageExtractionResult:
    ocr_text = _extract_ocr_text_from_page(page, abs_path)
    direct_text = _extract_text_pdf_page(page)
    return _select_page_text(direct_text=direct_text, ocr_text=ocr_text)


def _collect_pdf_page_snapshots(abs_path: Path) -> tuple[list[PageTextSnapshot], str]:
    try:
        import fitz
    except ImportError:
        logger.warning("PyMuPDF is not installed. Skipping PDF extraction for %s", abs_path)
        return [], "unsupported"

    page_snapshots: list[PageTextSnapshot] = []
    document = fitz.open(abs_path)

    try:
        document_signal = _classify_pdf_document(document)
        logger.info(
            "PDF classification file=%s kind=%s sample_scores=%s",
            abs_path,
            document_signal.source_type,
            [signal.direct_score for signal in document_signal.page_signals],
        )

        for page in document:
            if document_signal.source_type == "digital_pdf_text":
                direct_text = _extract_text_pdf_page(page)
                page_result = PageExtractionResult(
                    selected_text=direct_text,
                    direct_text=direct_text,
                    ocr_text=None,
                    direct_score=score_extracted_text(direct_text),
                    ocr_score=0,
                    used_ocr=False,
                )
            elif document_signal.source_type == "scanned_pdf_ocr":
                page_result = _extract_scanned_pdf_page(page, abs_path)
            else:
                page_result = _extract_mixed_pdf_page(page, abs_path)

            page_snapshots.append(
                PageTextSnapshot(
                    page_number=page.number + 1,
                    direct_text=page_result.direct_text,
                    selected_text=page_result.selected_text,
                    source_mode="ocr" if page_result.used_ocr else "direct",
                )
            )
    finally:
        document.close()

    return page_snapshots, document_signal.source_type


def _finalize_extracted_text_result(
    *,
    abs_path: Path,
    source_type: str,
    raw_text: str | None,
    normalized_text: str | None,
) -> ExtractedTextResult:
    quality_score = compute_quality_score(normalized_text)
    document_kind = infer_document_kind(
        abs_path=abs_path,
        normalized_text=normalized_text,
        source_type=source_type,
    )
    llm_cleaning_used = False
    llm_reason: str | None = None
    llm_error: str | None = None

    should_run_llm, llm_reason = should_run_llm_cleaning(
        source_type=source_type,
        document_kind=document_kind,
        normalized_text=normalized_text,
    )
    if should_run_llm and normalized_text:
        llm_output, llm_error = run_llm_normalization(
            source_type=source_type,
            document_kind=document_kind,
            normalized_text=normalized_text,
        )
        if is_llm_output_acceptable(source_text=normalized_text, candidate_text=llm_output):
            normalized_text = llm_output
            quality_score = compute_quality_score(normalized_text)
            llm_cleaning_used = True
        elif llm_error:
            logger.warning(
                "LLM normalization skipped after failure file=%s source_type=%s document_type=%s error=%s",
                abs_path,
                source_type,
                document_kind,
                llm_error,
            )

    extract_strategy = build_strategy(source_type, document_kind, quality_score)
    if llm_cleaning_used:
        extract_strategy = f"{extract_strategy}_llm"

    return ExtractedTextResult(
        extracted_text=normalized_text,
        extract_status="SUCCESS" if normalized_text else "FAILED",
        extract_strategy=extract_strategy,
        extract_quality_score=quality_score,
        source_type=source_type,
        document_type=document_kind,
        raw_text=raw_text,
        normalized_text=normalized_text,
        extract_meta={
            "source_type": source_type,
            "document_type": document_kind,
            "extractor": (
                "pymupdf"
                if source_type == "digital_pdf_text"
                else "pymupdf+rapidocr"
                if source_type in {"scanned_pdf_ocr", "mixed_pdf"}
                else "python-docx"
                if source_type == "docx_text"
                else "plain-text"
            ),
            "llm_cleaning": llm_cleaning_used,
            "llm_reason": llm_reason,
            "llm_error": llm_error,
            "quality_score": quality_score,
            "strategy": extract_strategy,
        },
    )


def _extract_pdf_text(abs_path: Path) -> ExtractedTextResult:
    page_snapshots, source_type = _collect_pdf_page_snapshots(abs_path)
    raw_text, normalized_text = normalize_document_pages(page_snapshots, source_type=source_type)
    return _finalize_extracted_text_result(
        abs_path=abs_path,
        source_type=source_type,
        raw_text=raw_text,
        normalized_text=normalized_text,
    )


def _extract_text_file(abs_path: Path) -> ExtractedTextResult:
    raw_text = _extract_plain_text(abs_path)
    normalized_text = normalize_extracted_text(raw_text or "")
    return _finalize_extracted_text_result(
        abs_path=abs_path,
        source_type="plain_text",
        raw_text=raw_text,
        normalized_text=normalized_text,
    )


def _extract_docx_file(abs_path: Path) -> ExtractedTextResult:
    raw_text = _extract_docx_text(abs_path)
    normalized_text = normalize_extracted_text(raw_text or "")
    return _finalize_extracted_text_result(
        abs_path=abs_path,
        source_type="docx_text",
        raw_text=raw_text,
        normalized_text=normalized_text,
    )


def extract_text_from_file(file_path: str, file_ext: str | None) -> ExtractedTextResult:
    try:
        abs_path = resolve_absolute_path(file_path)
        extension = (file_ext or abs_path.suffix.removeprefix(".")).lower()

        if extension == "txt":
            result = _extract_text_file(abs_path)
        elif extension == "pdf":
            result = _extract_pdf_text(abs_path)
        elif extension == "docx":
            result = _extract_docx_file(abs_path)
        else:
            result = ExtractedTextResult(
                extracted_text=None,
                extract_status="FAILED",
                extract_strategy="unsupported",
                source_type="unsupported",
                document_type="unknown",
            )

        logger.info(
            "Extraction finished file=%s strategy=%s source_type=%s document_type=%s quality=%.4f status=%s",
            abs_path,
            result.extract_strategy,
            result.source_type,
            result.document_type,
            result.extract_quality_score,
            result.extract_status,
        )
        return result
    except Exception:
        logger.exception(
            "Text extraction failed for file_path=%s file_ext=%s",
            file_path,
            file_ext,
        )
        return ExtractedTextResult(
            extracted_text=None,
            extract_status="FAILED",
            extract_strategy="failed",
            source_type="unknown",
            document_type="unknown",
        )

